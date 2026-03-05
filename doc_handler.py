"""
doc_handler.py
──────────────
Handles all .docx file operations:
  • Extracting structured content (paragraphs, headings, list items)
  • Rebuilding a new .docx from corrected content blocks

Each block is a dict with:
    {
        "type":  "heading" | "list_item" | "paragraph",
        "text":  "<raw text>",
        "level": <int>          # only for headings (1–9)
    }
"""

from __future__ import annotations

import io
from typing import BinaryIO

from docx import Document
from docx.shared import Pt


# ── Style‑name constants used by python‑docx ──────────────────────────────────
_HEADING_PREFIX = "Heading"
_LIST_STYLES = {"List Bullet", "List Number", "List Paragraph"}


def extract_content(file: BinaryIO) -> list[dict]:
    """
    Read a .docx file and return an ordered list of content blocks.

    Each block preserves enough metadata to reconstruct the document later
    while giving the AI service plain text to correct.
    """
    doc = Document(file)
    blocks: list[dict] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            # Preserve empty paragraphs as spacing markers
            blocks.append({"type": "paragraph", "text": ""})
            continue

        style_name = para.style.name if para.style else ""

        if style_name.startswith(_HEADING_PREFIX):
            # Extract heading level from style name, e.g. "Heading 2" → 2
            try:
                level = int(style_name.replace(_HEADING_PREFIX, "").strip())
            except ValueError:
                level = 1
            blocks.append({"type": "heading", "text": text, "level": level})

        elif style_name in _LIST_STYLES or _is_list_paragraph(para):
            blocks.append({"type": "list_item", "text": text})

        else:
            blocks.append({"type": "paragraph", "text": text})

    return blocks


def rebuild_document(text: str) -> io.BytesIO:
    """
    Build a new .docx from corrected text block and return it as a
    BytesIO buffer ready for streaming to the client.
    """
    doc = Document()

    # Apply a clean default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    for block in text.split("\n"):
        text_line = block.strip()
        if text_line:
            doc.add_paragraph(text_line, style="Normal")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ── Helpers ───────────────────────────────────────────────────────────────────

def _is_list_paragraph(para) -> bool:
    """
    Detect list items that don't have an explicit list style but use
    numbering (numPr) in their XML properties — common in many .docx files.
    """
    pPr = para._element.find(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr"
    )
    if pPr is None:
        return False
    numPr = pPr.find(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr"
    )
    return numPr is not None
